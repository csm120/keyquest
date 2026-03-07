from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt
import os
import subprocess


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    font_el = OxmlElement("w:rFonts")
    font_el.set(qn("w:ascii"), "Arial")
    font_el.set(qn("w:hAnsi"), "Arial")
    run_props.append(font_el)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), "24")
    run_props.append(font_size)

    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bold_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def resave_with_word(output_path):
    """If Microsoft Word is available, resave as the current default DOCX format."""
    if os.name != "nt":
        return

    ps_script = rf"""
$ErrorActionPreference = 'Stop'
$path = '{output_path.replace("'", "''")}'
$word = $null
$doc = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open($path)
    $wdFormatDocumentDefault = 16
    $doc.SaveAs([ref]$path, [ref]$wdFormatDocumentDefault)
}}
finally {{
    if ($doc -ne $null) {{ $doc.Close() }}
    if ($word -ne $null) {{ $word.Quit() }}
}}
"""
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps_script],
            check=False,
            capture_output=True,
            text=True,
        )
    except Exception:
        # Keep the DOCX even if Word automation is unavailable.
        pass


def build_post(output_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    for heading_level in range(1, 4):
        style_name = f"Heading {heading_level}"
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Arial"

    doc.add_heading(
        "KeyQuest: Let's Have Some Fun While Learning to Type", level=1
    )

    doc.add_paragraph(
        'I wanted a typing program that felt useful, accessible, and genuinely fun. '
        'Not endless mindless repeat typing drills. '
        'If you know, you know: "Anyone remember, Dad had a fall? Poor old Dad." '
        'Or, "The bed broke, what a classic!" '
        'That kind of repetition wears people out quickly.'
    )

    doc.add_paragraph(
        'So with all that out of the way, here is the short version: '
        'KeyQuest is built to help you improve your typing skills while keeping '
        'the experience engaging, practical, and a little more human. '
        'It has more personality than a drill sergeant, and yes, that is '
        'absolutely on purpose.'
    )

    doc.add_heading("What KeyQuest Is", level=2)
    doc.add_paragraph(
        'KeyQuest is an accessibility-first typing program for keyboard-only '
        'users, screen reader users, and anyone who wants to build real typing '
        'speed and accuracy without getting stuck in one boring loop. You get '
        'structured lessons, speed tests, sentence practice, games, progress '
        'tracking, and motivation systems that actually make you want to keep '
        'going.'
    )

    doc.add_paragraph(
        'This is not a program that teaches you the home row and then leaves '
        'you bored to tears! KeyQuest has 33 progressive lessons that cover all '
        '104 keys on a standard keyboard: letters, numbers, punctuation, '
        'function keys, and special keys. Each lesson introduces one or two '
        'new keys at a time, and the program adapts as you go. If you are '
        'doing well, lessons get shorter. If you are struggling, lessons '
        'extend to give you more practice. You never get stuck repeating the '
        'same thing over and over, and you never get pushed ahead before you '
        'are ready.'
    )

    doc.add_paragraph(
        'Letters are announced phonetically, so you hear '
        '"A as in Apple" or "B as in Bravo" instead of just the letter name. '
        'This makes a real difference when you are learning by ear.'
    )

    doc.add_heading("Who It May Help", level=2)
    doc.add_paragraph(
        'KeyQuest may help independent learners, people coming back to '
        'typing after a long break, screen reader users, low-vision users, '
        'and people who simply want a more human way to practice.'
    )

    doc.add_paragraph(
        'It can also be useful for assistive technology instructors, '
        'teachers, and trainers who want one program that can support '
        'keyboard basics, typing practice, shortcut review, and confidence '
        'building in the same place. Instead of jumping between multiple '
        'tools, they can use one program that covers a lot of ground.'
    )

    doc.add_heading("How It Fits Into Your Workflow", level=2)
    doc.add_paragraph(
        'A little lesson workflow can still be useful, and sometimes it is '
        'exactly what you need. The difference here is balance. KeyQuest '
        'blends structured lessons with practice modes and games so you can '
        'stay focused without feeling boxed in.'
    )

    doc.add_paragraph('Here is what you have to work with beyond the lessons:')

    p = doc.add_paragraph()
    p.add_run('Keyboard Explorer. ').bold = True
    p.add_run(
        'Press any key and hear its name, where it is on the keyboard, and '
        'what it is commonly used for. This is great for getting oriented, '
        'especially if you are new to a full-size keyboard.'
    )

    p = doc.add_paragraph()
    p.add_run('Free Practice. ').bold = True
    p.add_run(
        'Practice the keys you have already learned without affecting your '
        'lesson progress. No pressure, no tracking, just practice.'
    )

    p = doc.add_paragraph()
    p.add_run('Speed Tests. ').bold = True
    p.add_run(
        'Set your own time limit and get a full breakdown of your words per '
        'minute and accuracy. KeyQuest tracks both your actual speed and your '
        'correct speed, so you get a more honest picture of how you are doing. '
        'You can copy your results when you want to save or share them.'
    )

    p = doc.add_paragraph()
    p.add_run('Sentence Practice. ').bold = True
    p.add_run(
        'Type real sentences organized by topic. Topics include general '
        'English, Spanish, Windows commands, JAWS commands, NVDA commands, '
        'science facts, history, geography, math, literature, and vocabulary. '
        'This is where your typing skills meet real-world content. You can '
        'also edit existing sentence files or add your own, which is '
        'especially useful if you want more Spanish practice or want to build '
        'a custom practice set for your own needs.'
    )

    p = doc.add_paragraph()
    p.add_run('Learn Sounds. ').bold = True
    p.add_run(
        'There is also a Learn Sounds area where you can hear what the '
        'built-in sounds mean. That is helpful if you are new to the program '
        'and want to understand its audio cues before jumping into practice.'
    )

    p = doc.add_paragraph()
    p.add_run('Help and updates. ').bold = True
    p.add_run(
        'KeyQuest now includes a built-in user guide, a plain-language What\'s '
        'New page, and a Check for Updates option from the Main Menu. That '
        'makes it easier to learn the program, keep up with changes, and keep '
        'up without digging through release notes.'
    )

    doc.add_heading("Learn While You Type", level=2)
    doc.add_paragraph(
        'Most typing programs give you random words or meaningless letter '
        'combinations. KeyQuest takes a different approach. The sentences you '
        'practice are designed to teach you something useful while you build '
        'your speed and accuracy.'
    )

    doc.add_heading("Keyboard Shortcuts and Commands", level=3)
    doc.add_paragraph(
        'If you have ever wanted to learn common keyboard shortcuts, KeyQuest '
        'has sentence sets that walk you through them as you type. The Windows '
        'Commands topic covers everyday tasks like showing the desktop, '
        'switching between open windows, and locking your computer. You are '
        'not just reading a list. You are typing it out, which helps it '
        'stick.'
    )

    doc.add_paragraph(
        'For screen reader users, there are dedicated JAWS and NVDA command '
        'sets. The JAWS sentences include commands for reading from the '
        'current spot, hearing the time or date, and hearing the window '
        'title. The NVDA sentences include commands for opening the NVDA '
        'menu, hearing the window title, and saving your settings. Even if '
        'you already know these commands, typing them out helps build '
        'muscle memory.'
    )

    doc.add_heading("Interesting Facts and Educational Content", level=3)
    doc.add_paragraph(
        'KeyQuest also packs in educational content across several topics. '
        'The Science Facts sentences teach you things like how photosynthesis '
        'converts light energy into chemical energy, how the human body '
        'contains approximately thirty seven trillion cells, and how DNA '
        'stands for deoxyribonucleic acid. The History sentences cover events '
        'like the Declaration of Independence being signed in seventeen '
        'seventy six and the Roman Empire falling in four hundred seventy six '
        'AD. Geography sentences teach you that Mount Everest is the tallest '
        'mountain on Earth and that the Sahara Desert is the largest hot '
        'desert in the world.'
    )

    doc.add_paragraph(
        'There are also Math sentences that cover concepts like the '
        'Pythagorean theorem and prime numbers, Literature sentences with '
        'classic quotes like "To be or not to be, that is the question" and '
        '"Brevity is the soul of wit," and Vocabulary Building sentences that '
        'help you practice words and definitions in context. Even the general '
        'English sentences are full of interesting tech trivia, like the fact '
        'that the first computer mouse was made of wood and that the word '
        '"typewriter" can be typed using only the top row of keys.'
    )

    doc.add_heading("Spanish Language Practice", level=3)
    doc.add_paragraph(
        'If you are learning Spanish or want to practice typing in a second '
        'language, KeyQuest includes a full set of Spanish sentences covering '
        'the same kind of interesting content. This is a nice option if you '
        'want to build bilingual typing skills without switching to a separate '
        'program.'
    )

    doc.add_heading("Games", level=2)
    doc.add_paragraph(
        'Three built-in games give you a break from structured practice while '
        'still building your skills.'
    )

    p = doc.add_paragraph()
    p.add_run('Letter Fall. ').bold = True
    p.add_run(
        'Letters fall from the top of the screen and you type them before they '
        'hit the bottom. The speed picks up as your score increases, so it '
        'stays challenging.'
    )

    p = doc.add_paragraph()
    p.add_run('Word Typing. ').bold = True
    p.add_run(
        'A fast-paced 30-second challenge where you type as many words as you '
        'can. You get a speed and accuracy breakdown at the end.'
    )

    p = doc.add_paragraph()
    p.add_run('Hangman. ').bold = True
    p.add_run(
        'Classic word guessing with a twist. KeyQuest uses an offline '
        'dictionary with over 919,000 words and real definitions. After each '
        'round, you can see the word\'s definition, and the program offers to '
        'let you practice typing a sentence with that word. You can navigate '
        'your guessed letters with LEFT ARROW, RIGHT ARROW, HOME, and END, so '
        'you always know where you are. Word selection is weighted so you get '
        'common word lengths most of the time, with the occasional long word '
        'thrown in to keep things interesting.'
    )

    doc.add_heading("Progression and Motivation", level=2)
    doc.add_paragraph(
        'KeyQuest has a full progression system designed to keep you coming '
        'back. This is not just a star rating at the end of a lesson. It is a '
        'layered system that gives you something to work toward every time you '
        'sit down.'
    )

    p = doc.add_paragraph()
    p.add_run('XP and Levels. ').bold = True
    p.add_run(
        'You earn experience points for keystrokes, completed lessons, badges, '
        'quests, and more. There are 10 levels, from Keyboard Novice all the '
        'way up to Keyboard Legend.'
    )

    p = doc.add_paragraph()
    p.add_run('Badges. ').bold = True
    p.add_run(
        'Over 30 achievement badges track your milestones: First Steps, '
        'Perfectionist, Speed Demon, Week Warrior, Dedication Master, and '
        'many more. Badges are awarded automatically as you hit specific goals.'
    )

    p = doc.add_paragraph()
    p.add_run('Quests. ').bold = True
    p.add_run(
        'Longer-term goals like Home Row Master, Accuracy Expert, Marathon '
        'Runner, and Game Champion give you something to chase across multiple '
        'sessions.'
    )

    p = doc.add_paragraph()
    p.add_run('Daily Challenges. ').bold = True
    p.add_run(
        'Each day of the week has a different challenge type. Speed Monday, '
        'Accuracy Tuesday, Sentence Wednesday, Game Thursday, Focus Friday, '
        'and Marathon Weekend each push you in a different direction with bonus '
        'XP as the reward. Sunday is a rest day.'
    )

    p = doc.add_paragraph()
    p.add_run('Streaks. ').bold = True
    p.add_run(
        'KeyQuest tracks consecutive days of practice and encourages you to '
        'keep your streak alive.'
    )

    p = doc.add_paragraph()
    p.add_run('Virtual Pet. ').bold = True
    p.add_run(
        'Choose from six pet types: Robot, Dragon, Owl, Cat, Dog, or Phoenix. '
        'Your pet has five evolution stages and grows as you practice. Its mood '
        'changes based on your performance, and you can buy it accessories, '
        'toys, and food from the in-game shop. Yes, your pet gets happier when '
        'you keep typing and improving your skills, which makes progress feel '
        'a lot more rewarding.'
    )

    p = doc.add_paragraph()
    p.add_run('Shop. ').bold = True
    p.add_run(
        'Earn coins through gameplay and spend them on sound packs '
        '(Mechanical, Arcade, Sci-Fi, Nature, Musical), visual themes '
        '(Cyberpunk, Forest, Ocean, Retro, and more), pet accessories, and '
        'power-ups.'
    )

    doc.add_heading("Accessibility", level=2)
    doc.add_paragraph(
        'KeyQuest is accessibility-first by design. That means it was built '
        'from the ground up with screen reader users and keyboard-only users '
        'in mind, not bolted on after the fact.'
    )

    p = doc.add_paragraph()
    p.add_run('Screen reader integration. ').bold = True
    p.add_run(
        'KeyQuest works with JAWS, NVDA, and Narrator. If no screen reader is '
        'running, it uses built-in text-to-speech instead. You can set speech '
        'to auto-detect, screen reader only, text-to-speech only, or off.'
    )

    p = doc.add_paragraph()
    p.add_run('Keyboard-only navigation. ').bold = True
    p.add_run(
        'Everything in KeyQuest is navigable with the keyboard. Arrow keys, '
        'ENTER, SPACE, and ESCAPE are all you need. No mouse required.'
    )

    p = doc.add_paragraph()
    p.add_run('Escape safety. ').bold = True
    p.add_run(
        'Press ESCAPE three times from any active mode and you return to the '
        'main menu. This works consistently everywhere in the program, so you '
        'never get stuck.'
    )

    p = doc.add_paragraph()
    p.add_run('Theme support. ').bold = True
    p.add_run(
        'KeyQuest auto-detects your Windows theme and supports dark, light, '
        'and high contrast modes. Focus indicators and selection markers are '
        'always visible regardless of theme.'
    )

    p = doc.add_paragraph()
    p.add_run('Low-vision support. ').bold = True
    p.add_run(
        'Text can be made larger, up to 200 percent, and Focus Assist can add '
        'stronger emphasis to the part of the screen you should pay attention '
        'to. Longer prompts and typed text also wrap across multiple lines '
        'instead of being squeezed into one hard-to-read row.'
    )

    p = doc.add_paragraph()
    p.add_run('Visual feedback. ').bold = True
    p.add_run(
        'KeyQuest does not rely on speech alone. It can also show brief visual '
        'feedback for correct and incorrect key presses, clearer active-area '
        'grouping on typing screens, and an on-screen Escape counter while you '
        'leave active modes.'
    )

    p = doc.add_paragraph()
    p.add_run('Phonetic feedback. ').bold = True
    p.add_run(
        'Every letter is announced with its phonetic equivalent, making it '
        'easier to distinguish similar-sounding letters.'
    )

    doc.add_paragraph(
        'Anyone can use KeyQuest. It is accessibility-first by design, which '
        'often makes it clearer and more efficient for many users, not just '
        'those using assistive technology.'
    )

    doc.add_heading("Two Ways to Run KeyQuest", level=2)
    doc.add_paragraph(
        'You can run the program either way, depending on what you prefer:'
    )
    paragraph = doc.add_paragraph(style='List Number')
    add_hyperlink(
        paragraph,
        'Installer.exe',
        'https://github.com/csm120/KeyQuest/releases/latest/download/KeyQuestSetup.exe',
    )
    paragraph.add_run(
        ': Great if you want a standard setup with shortcuts and a familiar '
        'install flow.'
    )
    paragraph = doc.add_paragraph(style='List Number')
    add_hyperlink(
        paragraph,
        'Portable.zip',
        'https://github.com/csm120/KeyQuest/releases/latest/download/KeyQuest-win64.zip',
    )
    paragraph.add_run(
        ': Great if you prefer to extract and run KeyQuest.exe directly '
        'without a full install.'
    )

    doc.add_paragraph(
        'Both the installer and portable version can update from inside the '
        'program. KeyQuest also keeps your progress during updates, and if you '
        'added your own sentence files, it keeps those too while still bringing '
        'in new built-in content.'
    )

    paragraph = doc.add_paragraph('If you want the full user guide, open the ')
    add_hyperlink(
        paragraph,
        'KeyQuest User Guide',
        'https://csm120.github.io/KeyQuest/',
    )
    paragraph.add_run('.')

    paragraph = doc.add_paragraph('If you want the latest plain-language changes, open ')
    add_hyperlink(
        paragraph,
        'New in Key Quest',
        'https://csm120.github.io/KeyQuest/changelog.html',
    )
    paragraph.add_run('.')

    paragraph = doc.add_paragraph(
        'If you want the installer route, grab the '
    )
    add_hyperlink(
        paragraph,
        'KeyQuest Installer.exe',
        'https://github.com/csm120/KeyQuest/releases/latest/download/KeyQuestSetup.exe',
    )
    paragraph.add_run('.')

    paragraph = doc.add_paragraph(
        'If you want the portable route, download the '
    )
    add_hyperlink(
        paragraph,
        'KeyQuest Portable.zip',
        'https://github.com/csm120/KeyQuest/releases/latest/download/KeyQuest-win64.zip',
    )
    paragraph.add_run('.')

    paragraph = doc.add_paragraph(
        'If you want to see the source code, visit the '
    )
    add_hyperlink(
        paragraph,
        'KeyQuest GitHub repository',
        'https://github.com/csm120/KeyQuest',
    )
    paragraph.add_run('.')

    doc.add_heading("FAQ", level=2)

    add_bold_paragraph(doc, 'Is KeyQuest only for blind users?')
    doc.add_paragraph(
        'No. Anyone can use it. It is accessibility-first by design, which '
        'often makes it clearer and more efficient for many users.'
    )

    add_bold_paragraph(doc, 'Does KeyQuest work with screen readers?')
    doc.add_paragraph(
        'Yes. KeyQuest works with JAWS, NVDA, and Narrator. If no screen '
        'reader is detected, it uses built-in text-to-speech as a fallback.'
    )

    add_bold_paragraph(doc, 'Is it only speech-based?')
    doc.add_paragraph(
        'No. It also includes visual feedback, larger text options, strong '
        'focus cues, and support for Windows High Contrast. That makes it '
        'useful for many low-vision users and for people who want both audio '
        'and visual guidance.'
    )

    add_bold_paragraph(doc, 'What keys does KeyQuest teach?')
    doc.add_paragraph(
        'All 104 keys on a standard keyboard. That includes letters, numbers, '
        'punctuation, function keys, and special keys like INSERT, DELETE, '
        'HOME, END, and the arrow keys.'
    )

    add_bold_paragraph(doc, 'Does it only do basic drills?')
    doc.add_paragraph(
        'No. The whole point is to avoid endless mindless repeat typing '
        'drills. You get adaptive lessons, multiple practice modes, three '
        'games, a progression system, and a virtual pet. Structure with '
        'variety.'
    )

    add_bold_paragraph(doc, 'Can I learn keyboard shortcuts while I type?')
    doc.add_paragraph(
        'Yes. KeyQuest includes sentence practice sets for Windows commands, '
        'JAWS commands, and NVDA commands. You type out real shortcuts and '
        'commands, which helps you learn and remember them.'
    )

    add_bold_paragraph(doc, 'Does KeyQuest have educational content?')
    doc.add_paragraph(
        'Yes. Sentence practice topics include science facts, history, '
        'geography, math, literature quotes, vocabulary building, and Spanish '
        'language practice. You learn while you type.'
    )

    add_bold_paragraph(doc, 'Does KeyQuest have games?')
    doc.add_paragraph(
        'Yes. Three games: Letter Fall, Word Typing, and Hangman. Each one '
        'builds typing skills in a different way while keeping things fun.'
    )

    add_bold_paragraph(doc, 'Do I have to install it?')
    doc.add_paragraph(
        'No. You can install it with the EXE or run it portably from the ZIP. '
        'Pick the one that matches your setup style.'
    )

    add_bold_paragraph(doc, 'Can I track improvement over time?')
    doc.add_paragraph(
        'Yes. KeyQuest tracks speed and accuracy trends, awards badges for '
        'milestones, logs your daily streaks, and lets you review results '
        'session by session.'
    )

    add_bold_paragraph(doc, 'Could a teacher or AT instructor use this with learners?')
    doc.add_paragraph(
        'Yes. It is useful for guided instruction because it combines keyboard '
        'orientation, structured lessons, practice modes, shortcuts, and '
        'motivating progress systems in one place. That can make it easier to '
        'support beginners without switching between several different tools.'
    )

    add_bold_paragraph(doc, 'Is KeyQuest free?')
    doc.add_paragraph(
        'Yes. KeyQuest is free and open source under the MIT license.'
    )

    add_bold_paragraph(doc, 'Does KeyQuest save my progress?')
    doc.add_paragraph(
        'Yes. All progress is saved locally, including lesson completion, '
        'stars, badges, XP, pet evolution, shop purchases, and streak history.'
    )

    add_bold_paragraph(doc, 'How do I keep up with updates?')
    doc.add_paragraph(
        'You can use Check for Updates from inside the program, open the '
        'built-in guide, or read the plain-language What\'s New page online. '
        'That makes it easier to keep up.'
    )

    doc.add_heading("Final Thoughts", level=2)
    doc.add_paragraph(
        'If you have wanted a typing tool that is serious about growth but '
        'still enjoyable to come back to, KeyQuest is worth trying. Pick '
        'installer or portable, start where you are, and keep moving forward. '
        'This is the typing program I wish I had when I was learning to type '
        'myself. I hope you love it!'
    )
    final_para = doc.add_paragraph()
    final_para.add_run("You still here? KeyQuest won't run itself you know... Need the links again?")
    final_para = doc.add_paragraph()
    add_hyperlink(
        final_para,
        'Get KeyQuest Installer.exe',
        'https://github.com/csm120/KeyQuest/releases/latest/download/KeyQuestSetup.exe',
    )
    final_para = doc.add_paragraph()
    add_hyperlink(
        final_para,
        'Get KeyQuest Portable.zip',
        'https://github.com/csm120/KeyQuest/releases/latest/download/KeyQuest-win64.zip',
    )
    final_para = doc.add_paragraph()
    final_para.add_run('Qapla!')

    doc.save(output_path)
    resave_with_word(output_path)


if __name__ == "__main__":
    build_post(r"C:\OneDrive\Downloads\KeyQuest Lets Have Some Fun While Learning to Type.docx")
